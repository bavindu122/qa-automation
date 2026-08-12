"""Build the assignment report as a styled DOCX for deterministic PDF export."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "report"
OUTPUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUTPUT / "nopcommerce-qa-automation-report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(90, 99, 110)
INK = RGBColor(31, 45, 61)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_table_text(table, header=True, size=9):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
                    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
                    run.font.size = Pt(size)
                    if header and row_index == 0:
                        run.bold = True
                        run.font.color.rgb = INK
            if header and row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    if header:
        set_repeat_table_header(table.rows[0])


def add_table(doc, headers, rows, widths_dxa, size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    set_table_geometry(table, widths_dxa)
    style_table_text(table, size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead:
        paragraph.add_run(bold_lead).bold = True
    paragraph.add_run(text)
    return paragraph


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "UCSC SQA | nopCommerce QA Automation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def build_report():
    doc = Document()
    configure_document(doc)

    # editorial_cover pattern, using restrained business-report styling.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(110)
    kicker = doc.add_paragraph("PRACTICAL TAKE-HOME ASSIGNMENT")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.runs[0].bold = True
    kicker.runs[0].font.size = Pt(11)
    kicker.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph("nopCommerce QA Automation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(30)
    title.runs[0].font.color.rgb = INK
    subtitle = doc.add_paragraph("Software Quality Assurance Analysis, Test Design and Selenium Implementation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(70)
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = MUTED
    metadata = [
        ("Student", "Bavindu Shamen"),
        ("Application", "nopCommerce public demo store"),
        ("Technology", "Java 21 | Selenium 4 | TestNG 7 | Maven | Page Object Model"),
        ("Date", "12 August 2026"),
        ("Repository", "nopcommerce-qa-automation"),
    ]
    add_table(doc, ["Report detail", "Value"], metadata, [2700, 6660], size=10)
    note = doc.add_paragraph("Submission note: the public demo's Cloudflare verification blocked the recorded Selenium run. This is reported as an environment limitation, not a pass or product defect.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = MUTED

    doc.add_page_break()
    add_heading(doc, "Executive summary", 1)
    add_body(doc, "This project evaluates the customer-facing nopCommerce demo store and implements a maintainable browser-automation suite for five high-value regression scenarios. The work follows a traceable sequence: application selection, requirement analysis, manual scenario design, automation decisions, framework implementation, execution, and controlled debugging.")
    add_body(doc, "Eighteen manual scenarios cover registration, authentication, search, catalog, cart, comparison, and guest checkout. Five were automated using Java, Selenium WebDriver, TestNG, Maven, and Page Object Model. The project includes configuration overrides, explicit waits, isolated browser sessions, and failure screenshots.")
    add_body(doc, "The source compiled successfully and the controlled debugging test passed after its intentional assertion mismatch was corrected. A live full-suite attempt was blocked before application interaction by nopCommerce's Cloudflare security-verification page. The limitation is preserved transparently and the suite remains ready for an authorised rerun.")

    add_heading(doc, "1. Application selection and scope", 1)
    add_body(doc, "The selected system is the official public demonstration storefront at https://demo.nopcommerce.com/. It models realistic business-to-consumer journeys without requiring access to a production system or real financial information.")
    suitability = [
        ("Functional breadth", "Authentication, search, product detail, cart, comparison, and checkout support broad test design."),
        ("Automation fit", "Distinct pages and stable user-visible outcomes support reusable Page Objects."),
        ("Safety", "Synthetic customer and address data can be used in a public demo environment."),
        ("Positive/negative coverage", "Forms and searches provide successful, invalid, and empty-result behaviours."),
    ]
    add_table(doc, ["Criterion", "Assessment"], suitability, [2300, 7060])
    add_body(doc, "In scope are the desktop Chrome storefront, customer-visible validation, cart and comparison state, and guest checkout confirmation. Administration, database/API validation, real payment processing, email delivery, load/security testing, and pixel-perfect visual comparison are excluded.")

    add_heading(doc, "2. Requirement and risk analysis", 1)
    requirements = [
        ("FR-01", "Navigation", "Reach home, category, search, account, cart and comparison functions."),
        ("FR-02", "Registration", "Validate mandatory fields; accept unique users; reject duplicate email."),
        ("FR-03", "Authentication", "Accept valid credentials, reject invalid credentials, and support logout."),
        ("FR-04", "Discovery", "Search and browse products; explain empty results; support listing controls."),
        ("FR-05", "Product", "Show product details and enforce required configuration."),
        ("FR-06", "Cart", "Maintain products, quantities, prices, removal, and calculated totals."),
        ("FR-07", "Comparison", "Add, display, remove, and clear comparable products."),
        ("FR-08", "Checkout", "Validate terms and customer data; complete guest order confirmation."),
    ]
    add_table(doc, ["ID", "Area", "Expected behaviour"], requirements, [1000, 1700, 6660], size=8.7)
    risks = [
        ("Shared demo data and hourly reset", "High", "Create prerequisites per test; use guest and synthetic data."),
        ("Catalog/UI changes", "High", "Keep selectors in Page Objects and catalog values in configuration."),
        ("Asynchronous state", "High", "Use condition-based explicit waits; prohibit fixed sleeps."),
        ("Network/demo outage or challenge", "High", "Classify as environment failure and preserve evidence."),
        ("Browser/driver mismatch", "High", "Use Selenium Manager and record the execution versions."),
    ]
    add_table(doc, ["Risk", "Impact", "Mitigation"], risks, [2700, 1200, 5460], size=8.8)

    add_heading(doc, "3. Manual test design", 1)
    add_body(doc, "Priorities are risk based: P0 blocks a core journey or checkout; P1 materially affects discovery, account use, or cart integrity; P2 affects a supporting feature. Every scenario has explicit preconditions, data, steps, expected outcome, and evidence guidance in the repository's detailed manual-test document.")
    scenarios = [
        ("MTS-01", "Register with valid unique data", "P1", "FR-02"),
        ("MTS-02", "Validate required registration fields", "P1", "FR-02"),
        ("MTS-03", "Reject registration using an existing email", "P1", "FR-02"),
        ("MTS-04", "Log in with valid credentials", "P0", "FR-03"),
        ("MTS-05", "Reject invalid login credentials", "P0", "FR-03"),
        ("MTS-06", "Log out an authenticated customer", "P1", "FR-03"),
        ("MTS-07", "Search for an existing product", "P0", "FR-04"),
        ("MTS-08", "Search for a nonexistent product", "P1", "FR-04"),
        ("MTS-09", "Browse a catalog category", "P1", "FR-01/04"),
        ("MTS-10", "Sort a product listing by price", "P1", "FR-04"),
        ("MTS-11", "Add a simple product to the cart", "P0", "FR-05/06"),
        ("MTS-12", "Add a configured product to the cart", "P0", "FR-05/06"),
        ("MTS-13", "Update product quantity in the cart", "P0", "FR-06"),
        ("MTS-14", "Remove a product from the cart", "P1", "FR-06"),
        ("MTS-15", "Compare two products", "P2", "FR-07"),
        ("MTS-16", "Clear the product comparison list", "P2", "FR-07"),
        ("MTS-17", "Validate checkout terms requirement", "P0", "FR-08"),
        ("MTS-18", "Complete a valid guest checkout", "P0", "FR-08"),
    ]
    add_table(doc, ["ID", "Scenario", "Priority", "Trace"], scenarios, [1200, 5500, 1200, 1460], size=8.7)

    add_heading(doc, "4. Automation decision", 1)
    selected = [
        ("AUT-01", "MTS-05", "Invalid login", "Critical negative boundary; independent data and stable message."),
        ("AUT-02", "MTS-07", "Existing-product search", "Fast, frequent discovery path with relevance assertions."),
        ("AUT-03", "MTS-11", "Add simple product", "Critical catalog-to-cart state transition."),
        ("AUT-04", "MTS-15", "Compare two products", "Reusable repeated actions and cross-page state."),
        ("AUT-05", "MTS-18", "Guest checkout", "Revenue-critical end-to-end integration journey."),
    ]
    add_table(doc, ["Auto ID", "Manual", "Scenario", "Why selected"], selected, [1100, 1100, 2500, 4660], size=8.7)
    add_body(doc, "The other thirteen scenarios remain manual in this five-test scope because of persistent-account dependencies, duplicated mechanics, volatile pricing/configuration, lower business impact, or stronger exploratory value. MTS-13 quantity recalculation is the recommended next automation candidate.")

    add_heading(doc, "5. Automation framework", 1)
    architecture = [
        ("config", "Loads property defaults with command-line overrides."),
        ("driver", "Creates and closes WebDriver through Selenium Manager."),
        ("base", "Creates a clean browser before each test and always tears it down."),
        ("pages", "Owns private locators, explicit waits, and user actions."),
        ("models", "Represents immutable synthetic checkout data."),
        ("tests", "Expresses scenario intent and owns TestNG assertions."),
        ("listeners", "Captures timestamped screenshots when a test fails."),
    ]
    add_table(doc, ["Package", "Responsibility"], architecture, [2100, 7260], size=9.2)
    add_body(doc, "The design intentionally avoids direct element lookup in test classes, fixed sleeps, implicit waits, shared test order, and automatic retries. Mutable product values are held in config.properties. Each selected test receives a new browser session so cookies, cart items, and comparison state cannot leak between scenarios.")
    add_body(doc, "Locator preference is stable id, form name, component-scoped CSS, semantic relationship, then business text. ExpectedConditions wait for visibility, clickability, URL changes, selectable options, and user-visible notifications.")

    add_heading(doc, "6. Automated scenario implementation", 1)
    implementations = [
        ("AUT-01", "HomePage -> LoginPage", "Generic unsuccessful-login message is displayed."),
        ("AUT-02", "HomePage -> SearchResultsPage", "Query is retained and a relevant product is returned."),
        ("AUT-03", "ProductPage -> CartPage", "HTC smartphone exists in cart with quantity one."),
        ("AUT-04", "ProductPage -> CompareProductsPage", "HTC smartphone and Apple MacBook Pro are both displayed."),
        ("AUT-05", "Product -> Cart -> GuestCheckout -> Confirmation", "Success message and non-empty order number are displayed."),
    ]
    add_table(doc, ["ID", "Page flow", "Principal assertion"], implementations, [1100, 3400, 4860], size=8.8)
    add_body(doc, "Synthetic checkout data uses example.com and a non-real telephone number. The test chooses currently enabled shipping and payment methods, avoiding assumptions about a specific demo option. No real payment details are submitted.")

    add_heading(doc, "7. Execution and results", 1)
    environment = [
        ("Operating system", "macOS 26.5.1, Apple Silicon"),
        ("Java / Maven", "OpenJDK 21.0.12 / Maven 3.9.11"),
        ("Selenium / TestNG", "4.44.0 / 7.12.0"),
        ("Chrome / driver", "151.0.7922.76 / 151.0.7922.138"),
    ]
    add_table(doc, ["Environment item", "Recorded value"], environment, [2700, 6660], size=9.3)
    results = [
        ("Compile verification", "PASS", "Main and test sources compiled successfully."),
        ("Controlled debugging rerun", "PASS", "1 test; 0 failures, errors, or skips."),
        ("Five live Selenium scenarios", "BLOCKED", "Cloudflare verification replaced the app before first expected elements."),
    ]
    add_table(doc, ["Check", "Status", "Evidence"], results, [2800, 1300, 5260], size=9)
    add_body(doc, "The live suite result is not reported as a pass. Five screenshots show the same 'Performing security verification' page, and stack traces time out at each test's first nopCommerce locator. This common first divergence classifies the result as an external environment blocker. No attempt was made to bypass the site's security control.")

    doc.add_page_break()
    add_heading(doc, "8. Intentional failure and debugging", 1)
    add_body(doc, "Commit 36f3657 introduced a deterministic assertion failure. A captured generic login error contained 'Login was unsuccessful', while the temporary assertion searched for 'No customer account found'. TestNG reported one test, one failure, and an AssertionError at the assertion line.")
    debugging = [
        ("Observe", "Expected true but found false; failure isolated to the assertion."),
        ("Compare", "Known input and expected phrase did not describe the same message contract."),
        ("Classify", "Test defect, not application or runner defect."),
        ("Correct", "Assert the stable generic phrase: 'Login was unsuccessful'."),
        ("Verify", "Focused rerun completed with 1 test and 0 failures."),
    ]
    add_table(doc, ["Debug step", "Evidence and conclusion"], debugging, [1900, 7460], size=9.2)
    add_body(doc, "Commit dc535b2 preserves the correction and its explanation. Keeping the failure and fix as consecutive milestones creates reproducible evidence without leaving the current branch broken.")

    add_heading(doc, "9. Maintainability and limitations", 1)
    add_body(doc, "The Page Object boundary localises selector changes, system properties support environment overrides, and isolated sessions protect test independence. The repository's meaningful commit sequence allows the analysis, framework, scenarios, intentional failure, and fix to be reviewed independently.")
    add_body(doc, "The main limitation is dependency on a public shared demo whose catalog, reset cycle, availability, and bot protection are outside the project team's control. A dedicated authorised test environment would provide reliable CI execution and safer test-data lifecycle management. Until then, every execution should classify environment failures before interpreting application assertions.")

    add_heading(doc, "10. Conclusion and next steps", 1)
    add_body(doc, "The project fulfils the design and implementation objectives: a justified application choice, requirement analysis, 18 traceable manual scenarios, five risk-based automation decisions, a Java/Selenium/TestNG/POM framework, five automated flows, failure evidence, and a controlled debugging demonstration.")
    add_body(doc, "Before final submission, the student should rerun the suite when nopCommerce permits automation, update the execution record with the actual result, capture the required manual evidence, record the five-minute reflection in their own words, confirm lecturer registration of the selected demo if required, and insert the final GitHub and video links.")

    add_heading(doc, "Appendix A - Repository evidence", 1)
    commits = [
        ("566ec2c", "Initialize repository governance"),
        ("1bb6775", "Select nopCommerce demo application"),
        ("2ecc2c1", "Analyse requirements and risks"),
        ("7e28d04", "Define 18 manual scenarios"),
        ("0459016", "Select five automation scenarios"),
        ("cd9ca82", "Establish Selenium TestNG POM framework"),
        ("24d83f0", "Automate five regression scenarios"),
        ("36f3657", "Introduce controlled assertion failure"),
        ("dc535b2", "Correct assertion and document debugging"),
        ("9e05fa2", "Add video reflection and viva guides"),
    ]
    add_table(doc, ["Commit", "Milestone"], commits, [1900, 7460], size=9.2)

    add_heading(doc, "Appendix B - AI usage disclosure", 1)
    add_body(doc, "OpenAI Codex supported assignment decomposition, website comparison, live-flow review, implementation drafts, code review, debugging documentation, and report generation. The student remains responsible for reviewing and understanding all work, validating the live application's expected behaviour, executing the suite, explaining decisions during the viva, and maintaining the full activity log in AI_USAGE.md.")

    doc.core_properties.title = "nopCommerce QA Automation Assignment Report"
    doc.core_properties.author = "Bavindu Shamen"
    doc.core_properties.subject = "UCSC Software Quality Assurance practical take-home assignment"
    doc.core_properties.keywords = "SQA, Selenium, TestNG, nopCommerce, Page Object Model"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
